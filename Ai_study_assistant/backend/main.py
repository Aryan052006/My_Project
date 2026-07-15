from pydantic import BaseModel

from fastapi import FastAPI, UploadFile, File, Header
from fastapi.responses import FileResponse, StreamingResponse
import json
from fastapi.middleware.cors import CORSMiddleware
import os
import time

from pdf_utils import extract_text_from_pdf
from vector_store import (
    add_chunks,
    total_chunks,
    get_document_stats,
    document_exists,
    delete_document,
    rename_document,
    get_document_preview
)
from chunking import create_chunks
from embeddings import get_embedding
from semantic_search import find_top_k_chunks
from rag import generate_answer, generate_answer_stream, generate_followup_questions
from answer_sheet_service import generate_answer_sheet
from chat_memory import (
    add_message,
    get_history,
    clear_history,
    log_performance,
    get_analytics
)
from query_rewriter import rewrite_question
from question_parser import extract_marks
from settings_manager import load_settings, save_settings

app = FastAPI()

# Retrieval configurations are now dynamic from settings_manager

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


class ChatRequest(BaseModel):
    question: str
    session_id: str = "default"
    model: str = "qwen3:1.7b"


@app.get("/")
def home():
    return {
        "message": "AI Study Assistant Backend Running"
    }


@app.post("/chat")
def chat(request: ChatRequest, x_user_id: str = Header("default")):

    try:

        start_time = time.time()

        session_id = f"{x_user_id}_{request.session_id}"

        if total_chunks(user_id=x_user_id) == 0:

            return {
                "success": False,
                "message": "No PDF uploaded."
            }

        # =============================
        # Conversation History
        # =============================

        history = get_history(session_id)

        rewritten_question = rewrite_question(
            request.question,
            history
        )

        print("\n==============================")
        print("Original Question:")
        print(request.question)

        print("\nRewritten Question:")
        print(rewritten_question)
        print("==============================")

        # Save current user message AFTER rewriting
        add_message(
            session_id,
            "user",
            request.question
        )

        requested_marks = extract_marks(request.question)
        print(f"\nExtracted Marks: {requested_marks}")

        # =============================
        # Retrieval
        # =============================
        settings = load_settings()
        retrieval_start = time.time()

        results = find_top_k_chunks(
            rewritten_question,
            k=settings.get("top_k", 5),
            user_id=x_user_id
        )

        retrieval_end = time.time()

        print("\nRetrieved using:")
        print(rewritten_question)

        if len(results) == 0:

            return {
                "success": True,
                "answer": "Information not found in study material.",
                "sources": []
            }

        print(
            f"\nBest Similarity: {results[0]['score']:.4f}"
        )

        filtered_results = [
            result
            for result in results
            if result["score"] >= settings.get("min_similarity", 0.45)
        ]

        if len(filtered_results) == 0:

            return {
                "success": True,
                "answer": "Information not found in study material.",
                "sources": []
            }

        # =============================
        # Build Context
        # =============================

        contexts = [

            result["chunk"]["text"]

            for result in filtered_results

        ]

        sources = [

            {

                "pdf":
                result["chunk"]["source"],

                "chunk_id":
                result["chunk"]["id"],

                "similarity":
                round(
                    result["score"],
                    4
                ),

                "distance":
                round(
                    result["distance"],
                    4
                ),

                "preview":
                result["chunk"]["text"][:150]

            }

            for result in filtered_results

        ]

        # Get updated history (includes current user message)
        history = get_history(session_id)

        # =============================
        # Generate Answer
        # =============================

        generation_start = time.time()

        answer = generate_answer(
            question=request.question,
            contexts=contexts,
            history=history,
            marks=requested_marks,
            model=request.model
        )

        generation_end = time.time()

        # Save assistant reply
        add_message(
            session_id,
            "assistant",
            answer
        )

        # =============================
        # Logs
        # =============================

        print("\n========== PERFORMANCE ==========")

        print(
            f"Retrieval Time: {retrieval_end - retrieval_start:.2f} sec"
        )

        print(
            f"Generation Time: {generation_end - generation_start:.2f} sec"
        )

        print(
            f"Total Time: {generation_end - start_time:.2f} sec"
        )

        print("\nRetrieved Chunks")

        for result in filtered_results:

            print("=" * 60)

            print(
                "Similarity:",
                round(result["score"], 4)
            )

            print(
                "Distance:",
                round(result["distance"], 4)
            )

            print(
                "Source:",
                result["chunk"]["source"]
            )

            print(
                result["chunk"]["text"][:200]
            )

        print("=================================\n")

        return {

            "success": True,

            "answer": answer,

            "sources": sources

        }

    except Exception as e:

        import traceback

        traceback.print_exc()

        return {

            "success": False,

            "error": str(e)

        }
        
@app.post("/chat-stream")
async def chat_stream_endpoint(request: ChatRequest, x_user_id: str = Header("default")):
    session_id = f"{x_user_id}_{request.session_id}"

    if total_chunks(user_id=x_user_id) == 0:
        async def err_gen():
            yield json.dumps({"type": "error", "message": "No PDF uploaded."}) + "\n"
        return StreamingResponse(err_gen(), media_type="application/x-ndjson")

    history = get_history(session_id)
    rewritten_question = rewrite_question(request.question, history)
    add_message(session_id, "user", request.question)
    requested_marks = extract_marks(request.question)

    settings = load_settings()
    retrieval_start = time.time()
    results = find_top_k_chunks(rewritten_question, k=settings.get("top_k", 5), user_id=x_user_id)
    retrieval_end = time.time()
    log_performance('retrieval', (retrieval_end - retrieval_start) * 1000)

    filtered_results = [r for r in results if r["score"] >= settings.get("min_similarity", 0.45)]
    
    if len(filtered_results) == 0:
        async def not_found_gen():
            yield json.dumps({
                "type": "metadata",
                "sources": [],
                "confidence": 0,
                "success": True
            }) + "\n"
            msg = "Information not found in study material."
            yield json.dumps({"type": "text", "content": msg}) + "\n"
            add_message(session_id, "assistant", msg)
        return StreamingResponse(not_found_gen(), media_type="application/x-ndjson")

    contexts = [r["chunk"]["text"] for r in filtered_results]
    
    # Calculate average confidence of top results (normalized score was 0.0 to 1.0)
    avg_confidence = int(sum(r["score"] for r in filtered_results) / len(filtered_results) * 100)

    sources = [{
        "pdf": r["chunk"]["source"],
        "chunk_id": r["chunk"]["id"],
        "similarity": round(r["score"], 4),
        "preview": r["chunk"]["text"][:150]
    } for r in filtered_results]

    updated_history = get_history(session_id)

    async def event_generator():
        # Yield metadata first (sources and confidence)
        yield json.dumps({
            "type": "metadata",
            "sources": sources,
            "confidence": avg_confidence,
            "success": True
        }) + "\n"

        full_answer = ""
        generation_start = time.time()
        # Yield text chunks
        for text_chunk in generate_answer_stream(request.question, contexts, updated_history, requested_marks, request.model):
            full_answer += text_chunk
            yield json.dumps({"type": "text", "content": text_chunk}) + "\n"
        
        generation_end = time.time()
        log_performance('generation', (generation_end - generation_start) * 1000)

        # Generate and yield followups
        followups = generate_followup_questions(full_answer, request.model)
        if followups:
            yield json.dumps({"type": "followups", "content": followups}) + "\n"

        # Save to DB
        add_message(session_id, "assistant", full_answer)

    return StreamingResponse(event_generator(), media_type="application/x-ndjson")

from fastapi import BackgroundTasks
from typing import Dict

# Simple in-memory tracker for background uploads
UPLOAD_STATUS: Dict[str, str] = {}

def process_pdf_background(file_path: str, filename: str, x_user_id: str):
    job_id = f"{x_user_id}_{filename}"
    try:
        UPLOAD_STATUS[job_id] = "processing"
        
        text = extract_text_from_pdf(file_path)
        chunks = create_chunks(text, source=filename)
        
        # Batch embeddings to avoid rate limits and increase speed
        from embeddings import get_embeddings_batch
        texts_to_embed = [chunk["text"] for chunk in chunks]
        
        # Batch size of 90 is safe for both Cohere (96 max) and Gemini (100 max)
        batch_size = 90
        all_embeddings = []
        for i in range(0, len(texts_to_embed), batch_size):
            batch_texts = texts_to_embed[i:i + batch_size]
            batch_embeddings = get_embeddings_batch(batch_texts)
            all_embeddings.extend(batch_embeddings)
            
        for chunk, embedding in zip(chunks, all_embeddings):
            chunk["embedding"] = embedding
            
        add_chunks(chunks, user_id=x_user_id)
        
        UPLOAD_STATUS[job_id] = "ready"
    except Exception as e:
        import traceback
        traceback.print_exc()
        UPLOAD_STATUS[job_id] = "failed"

@app.get("/upload-status")
def get_upload_status(x_user_id: str = Header("default")):
    user_prefix = f"{x_user_id}_"
    statuses = {k.replace(user_prefix, ""): v for k, v in UPLOAD_STATUS.items() if k.startswith(user_prefix)}
    return {"success": True, "statuses": statuses}

@app.post("/upload-pdf")
async def upload_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...), x_user_id: str = Header("default")):

    print("========== UPLOAD START ==========")

    try:

        print("1")

        file_path = os.path.join(
            UPLOAD_FOLDER,
            f"{x_user_id}_{file.filename}"
        )

        print("2")

        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        # Trigger background task
        background_tasks.add_task(process_pdf_background, file_path, file.filename, x_user_id)
        
        # Mark as processing immediately so the frontend knows
        UPLOAD_STATUS[f"{x_user_id}_{file.filename}"] = "processing"

        response = {
            "success": True,
            "filename": file.filename,
            "message": "Processing in background"
        }

        print("========== UPLOAD QUEUED ==========")
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

        return {
            "success": False,
            "error": str(e)
        }

@app.get("/documents")
def documents(x_user_id: str = Header("default")):

    stats = get_document_stats(user_id=x_user_id)

    documents = []

    for name, chunks in stats.items():

        documents.append(
            {
                "name": name,
                "chunks": chunks
            }
        )

    return {

        "success": True,

        "total_documents": len(documents),

        "documents": documents

    }

from question_extractor import extract_questions

@app.post("/extract-questions")
async def extract_questions_endpoint(file: UploadFile = File(...), x_user_id: str = Header("default")):
    try:
        file_path = os.path.join(UPLOAD_FOLDER, f"{x_user_id}_{file.filename}")
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        text = extract_text_from_pdf(file_path)
        questions = extract_questions(text)
        
        parsed = []
        for i, q in enumerate(questions):
            marks = extract_marks(q)
            parsed.append({
                "id": i,
                "text": q,
                "marks": marks
            })
            
        return {"success": True, "questions": parsed, "filename": file.filename}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

class GenerateAnswerRequest(BaseModel):
    question: str
    marks: int

@app.post("/generate-individual-answer")
def generate_individual_answer(req: GenerateAnswerRequest, x_user_id: str = Header("default")):
    try:
        from answer_generator import answer_question
        ans = answer_question(req.question, req.marks, user_id=x_user_id)
        return {"success": True, "answer": ans}
    except Exception as e:
        return {"success": False, "error": str(e)}

class ExportAnswersRequest(BaseModel):
    qa_pairs: list
    format: str

@app.post("/export-answers")
def export_answers(req: ExportAnswersRequest, x_user_id: str = Header("default")):
    try:
        import os
        os.makedirs("uploads", exist_ok=True)
        filename = f"uploads/{x_user_id}_generated_answers.{req.format}"
        if req.format == "pdf":
            from pdf_generator import generate_pdf
            questions = [p["question"] for p in req.qa_pairs]
            answers = [p["answer"] for p in req.qa_pairs]
            generate_pdf(questions, answers, filename=filename)
        elif req.format == "docx":
            from docx import Document
            doc = Document()
            doc.add_heading("Generated Answers", 0)
            for i, p in enumerate(req.qa_pairs):
                doc.add_heading(f"Q{i+1}: {p['question']}", level=1)
                doc.add_paragraph(p['answer'])
            doc.save(filename)
        elif req.format == "md":
            with open(filename, "w", encoding="utf-8") as f:
                f.write("# Generated Answers\n\n")
                for i, p in enumerate(req.qa_pairs):
                    f.write(f"### Q{i+1}: {p['question']}\n\n")
                    f.write(f"{p['answer']}\n\n---\n\n")
        else:
            return {"success": False, "error": "Unsupported format"}
            
        return {"success": True, "download_url": f"/download-file?file={filename}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

class TutorRequest(BaseModel):
    topic: str
    num_questions: int = 5
    difficulty: str = "Medium"
    taxonomy: str = "Understanding"

@app.post("/tutor/summary")
def tutor_summary(req: TutorRequest, x_user_id: str = Header("default")):
    try:
        from tutor import generate_summary
        summary = generate_summary(req.topic, req.difficulty, req.taxonomy, user_id=x_user_id)
        return {"success": True, "summary": summary}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/tutor/quiz")
def tutor_quiz(req: TutorRequest, x_user_id: str = Header("default")):
    try:
        from tutor import generate_quiz
        quiz = generate_quiz(req.topic, req.num_questions, req.difficulty, req.taxonomy, user_id=x_user_id)
        if isinstance(quiz, dict) and "error" in quiz:
            return {"success": False, "error": quiz["error"]}
        return {"success": True, "quiz": quiz}
    except Exception as e:
        return {"success": False, "error": str(e)}

class RevisionRequest(BaseModel):
    topic: str
    num_cards: int = 10

@app.post("/revision/flashcards")
def revision_flashcards(req: RevisionRequest, x_user_id: str = Header("default")):
    try:
        from revision import generate_flashcards
        cards = generate_flashcards(req.topic, req.num_cards, user_id=x_user_id)
        if isinstance(cards, dict) and "error" in cards:
            return {"success": False, "error": cards["error"]}
        return {"success": True, "flashcards": cards}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/revision/cheatsheet")
def revision_cheatsheet(req: RevisionRequest, x_user_id: str = Header("default")):
    try:
        from revision import generate_cheat_sheet
        sheet = generate_cheat_sheet(req.topic, user_id=x_user_id)
        return {"success": True, "cheatsheet": sheet}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/download-file")
def download_file(file: str):
    import os
    if os.path.exists(file):
        return FileResponse(path=file, filename=file)
    return {"success": False, "error": "File not found"}

@app.delete(
    "/documents/{filename}"
)
def remove_document(
    filename: str,
    x_user_id: str = Header("default")
):

    deleted = delete_document(
        filename, user_id=x_user_id
    )

    if not deleted:

        return {
            "success": False,
            "message": "Document not found."
        }

    return {
        "success": True,
        "deleted": filename
    }

class RenameRequest(BaseModel):
    new_name: str

@app.put("/documents/{filename}/rename")
def rename_doc(filename: str, request: RenameRequest, x_user_id: str = Header("default")):
    renamed = rename_document(filename, request.new_name, user_id=x_user_id)
    if not renamed:
        return {"success": False, "message": "Document not found."}
    
    # Also rename the physical file if it exists
    old_path = os.path.join(UPLOAD_FOLDER, f"{x_user_id}_{filename}")
    new_path = os.path.join(UPLOAD_FOLDER, f"{x_user_id}_{request.new_name}")
    if os.path.exists(old_path) and not os.path.exists(new_path):
        os.rename(old_path, new_path)
        
    return {"success": True, "renamed_to": request.new_name}

@app.post("/documents/{filename}/reindex")
def reindex_doc(filename: str, x_user_id: str = Header("default")):
    file_path = os.path.join(UPLOAD_FOLDER, f"{x_user_id}_{filename}")
    if not os.path.exists(file_path):
        return {"success": False, "message": "Original PDF not found on server."}
        
    delete_document(filename, user_id=x_user_id)
    text = extract_text_from_pdf(file_path)
    chunks = create_chunks(text, source=filename)
    for chunk in chunks:
        chunk["embedding"] = get_embedding(chunk["text"])
    add_chunks(chunks, user_id=x_user_id)
    
    return {"success": True, "new_chunks": len(chunks)}

@app.get("/documents/{filename}/preview")
def preview_doc(filename: str, x_user_id: str = Header("default")):
    preview_text = get_document_preview(filename, user_id=x_user_id)
    if preview_text:
        return {"success": True, "preview": preview_text}
    return {"success": False, "message": "No content found."}

@app.post("/clear-chat")
def clear_chat(session_id: str = "default", x_user_id: str = Header("default")):

    clear_history(
        f"{x_user_id}_{session_id}"
    )

    return {
        "success": True,
        "message": "Chat cleared."
    }

@app.get("/sessions")
def get_sessions(x_user_id: str = Header("default")):
    from chat_memory import get_all_sessions
    return {"sessions": get_all_sessions(x_user_id)}

@app.get("/chat/history")
def get_chat_history(session_id: str = "default", x_user_id: str = Header("default")):
    from chat_memory import get_all_history
    return {"history": get_all_history(f"{x_user_id}_{session_id}")}

@app.get("/analytics")
def analytics_endpoint(x_user_id: str = Header("default")):
    stats = get_document_stats(user_id=x_user_id)
    total_docs = len(stats)
    total_chunks = sum(stats.values())
    
    db_analytics = get_analytics(user_id=x_user_id)
    
    return {
        "success": True,
        "total_documents": total_docs,
        "total_chunks": total_chunks,
        "total_questions": db_analytics["total_questions"],
        "total_sessions": db_analytics["total_sessions"],
        "avg_retrieval_ms": db_analytics["avg_retrieval_ms"],
        "avg_generation_ms": db_analytics["avg_generation_ms"]
    }

@app.get("/models")
def get_models():
    # Since we moved to Groq, we return Groq supported models.
    return {"models": ["llama-3.1-8b-instant", "llama-3.3-70b-versatile", "mixtral-8x7b-32768"]}

@app.get("/settings")
def get_settings():
    return {"success": True, "settings": load_settings()}

class SettingsUpdate(BaseModel):
    chunk_size: int
    chunk_overlap: int
    top_k: int
    min_similarity: float
    temperature: float
    max_tokens: int

@app.post("/settings")
def update_settings(request: SettingsUpdate):
    settings = request.dict()
    save_settings(settings)
    return {"success": True, "settings": settings}